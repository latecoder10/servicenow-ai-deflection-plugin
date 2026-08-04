"""
Generate: AI Service Desk Knowledge Intelligence Platform — Requirements & Technical Specification
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os


def setup_page(doc, size="A4"):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(3.18)
    section.orientation = WD_ORIENTATION.PORTRAIT


def tune_styles(doc):
    body = doc.styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(6)

    for n, size, color in [(1, 18, 0x1F3A5F), (2, 14, 0x1F3A5F), (3, 12, 0x2E4057)]:
        s = doc.styles[f"Heading {n}"]
        s.font.name = "Calibri Light"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)
        s.paragraph_format.space_before = Pt(14 - 2 * n)
        s.paragraph_format.space_after = Pt(4)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph("AI Service Desk Knowledge Intelligence Platform", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph("Enterprise Pre-Incident Deflection & ServiceNow Knowledge RAG Technical Specification", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph("Version 2.5.0-SNAPSHOT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    p = doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph("CONFIDENTIAL — ENTERPRISE ARCHITECTURE")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p = doc.add_paragraph("Prepared by: AI & ITSM Platform Engineering Team")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def add_toc(doc):
    p = doc.add_paragraph("Table of Contents", style="Heading 1")
    p2 = doc.add_paragraph()
    run = p2.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click and choose Update Field to populate."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    for x in (fldChar1, instrText, fldChar2, fldChar3, fldChar4):
        run._r.append(x)


def add_table(doc, header, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, name in enumerate(header):
        hdr[i].text = name
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cells[c_idx].text = str(value)
            for p in cells[c_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    setup_page(doc)
    tune_styles(doc)

    # ── Cover Page ──
    add_cover(doc)
    doc.add_page_break()

    # ── Table of Contents ──
    add_toc(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. Project Overview & Business ROI Goals", level=1)

    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "The AI Service Desk Knowledge Intelligence Platform is an enterprise-grade Java 21 / Spring Boot 3.5+ "
        "pre-incident resolution and knowledge management engine. Sitting directly between self-service portals "
        "and ServiceNow ITSM, it executes real-time Retrieval-Augmented Generation (RAG) to intercept user issues "
        "before an incident ticket is submitted."
    )

    doc.add_heading("1.2 Scope & Business Objectives", level=2)
    doc.add_paragraph(
        "By combining Pinecone vector similarity search, Google Gemini 3.6 Flash synthesis, cross-encoder reranking, "
        "and multi-factor confidence scoring, the platform resolves IT tickets autonomously:"
    )
    goals = [
        "Pre-Ticket Interception: Provide step-by-step verified resolutions within < 1.2 seconds of typing.",
        "Cost Savings: Average savings of $15.50 per deflected incident compared to Tier-1 helpdesk handling.",
        "Target Deflection Rate: Achieve 45% - 65% pre-ticket resolution rate across standard enterprise categories (VPN, Passwords, MFA, SaaS access).",
        "Zero-Downtime Resilience: Resilience4j Circuit Breaker ensures unhandled failures gracefully fallback to creating ServiceNow incidents.",
        "Continuous Knowledge Sync: Automated background synchronization across ServiceNow KB, Confluence, and Runbooks."
    ]
    for goal in goals:
        doc.add_paragraph(goal, style="List Bullet")

    # ══════════════════════════════════════════════════════════════
    # 2. TECH STACK
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. Technology Stack & Multi-Module Architecture", level=1)

    doc.add_heading("2.1 Core Stack Specifications", level=2)
    add_table(doc,
        ["Component", "Technology", "Version", "Purpose"],
        [
            ["Language", "Java", "21 LTS", "Core runtime (Virtual Threads, Pattern Matching, Records)"],
            ["Framework", "Spring Boot", "3.4.2 / 3.5.0", "Enterprise REST API, Dependency Injection, CQRS"],
            ["AI Framework", "Spring AI", "1.0.0-M6", "Standardized LLM, Embedding, and Reranking abstractions"],
            ["Vector DB", "Pinecone", "768-dim Text-Embedding-004", "High-performance vector similarity search index"],
            ["LLM Engine", "Google Gemini 3.6 Flash", "Google GenAI API", "RAG context synthesis & cross-encoder reranking"],
            ["ITSM Integration", "ServiceNow REST v2", "OAuth2 Bearer", "Incident ticket creation & KB synchronization"],
            ["Doc Ingestion", "Apache Tika", "2.9.2", "Multi-format parsing (PDF, Word, Excel, Confluence)"],
            ["Relational DB", "PostgreSQL", "16-alpine", "Audit trails, execution metrics, document metadata"],
            ["Cache & Limits", "Redis", "7-alpine", "Rate limiting token bucket & response caching"],
            ["Resilience", "Resilience4j", "2.2.0", "Circuit breaker, sliding window metrics, fallback queue"],
            ["Observability", "Prometheus & Grafana", "Latest", "Real-time deflection metrics & latency dashboards"],
        ]
    )

    doc.add_heading("2.2 Multi-Module Structure", level=2)
    doc.add_paragraph("The codebase is structured as a Maven multi-module Clean Architecture repository:")
    modules = [
        ("common", "Shared domain models, CorrelationContext, RFC 7807 ProblemDetails, base exceptions"),
        ("domain", "Core DDD entities (Incident, KnowledgeChunk, ResolutionSuggestion), Value Objects, Domain Ports"),
        ("application", "CQRS Use Cases (SuggestResolutionUseCase), Confidence Calculator, Deflection Engine"),
        ("knowledge-loader", "Apache Tika Document Parsers, SlidingWindowChunker (512 token chunks / 64 overlap)"),
        ("integration/pinecone", "Pinecone Vector Database Adapter for similarity search and vector indexing"),
        ("integration/servicenow", "ServiceNow REST v2 Client, OAuth2 authentication, Circuit Breaker integration"),
        ("integration/llm", "Spring AI Gemini 3.6 Flash Adapter, Text-Embedding-004 model integration"),
        ("analytics", "DeflectionMetrics calculation, ROI metric tracking, Micrometer telemetry"),
        ("security", "Spring Security, JJWT Bearer Token validation, Rate Limiting filter"),
        ("infrastructure", "AOP Audit Logger, Scheduled ServiceNow KB Sync, Resilience4j configuration"),
        ("api", "Spring REST Controllers, DTOs, OpenAPI 3.0 / Swagger UI, Spring Boot entry point")
    ]
    for mod, desc in modules:
        p = doc.add_paragraph()
        p.add_run(f"{mod}: ").bold = True
        p.add_run(desc)

    # ══════════════════════════════════════════════════════════════
    # 3. ARCHITECTURE APPROACH
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. Hexagonal Architecture & Design Patterns", level=1)

    doc.add_heading("3.1 Clean Hexagonal Architecture (Ports & Adapters)", level=2)
    doc.add_paragraph(
        "The system decouples core business logic from infrastructure using Hexagonal Architecture. "
        "The Domain layer defines inbound ports (e.g., SuggestResolutionUseCase) and outbound ports "
        "(VectorDatabasePort, ServiceNowPort, LlmPort). Infrastructure adapters implement these ports."
    )

    doc.add_heading("3.2 Applied Design Patterns", level=2)
    patterns = [
        ("Ports & Adapters", "Decouples LLM, Vector DB, and ServiceNow integrations from core resolution logic"),
        ("CQRS (Command Query Responsibility Segregation)", "Separates suggestion queries from document ingestion commands"),
        ("Strategy Pattern", "Allows dynamic switching between Pinecone and local in-memory vector indexing"),
        ("Circuit Breaker & Fallback", "Resilience4j wraps ServiceNow calls to buffer incidents locally if API is offline"),
        ("AOP Aspect-Oriented Logging", "AuditLogAspect wraps application methods to capture execution metrics and correlation IDs")
    ]
    for name, desc in patterns:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    # ══════════════════════════════════════════════════════════════
    # 4. LOW-LEVEL DESIGN (LLD)
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. Low-Level Design (LLD)", level=1)

    doc.add_heading("4.1 Pre-Incident Deflection Pipeline", level=2)
    steps = [
        "Self-Service Portal sends query title, description, user email, and department to POST /api/v1/suggestions/resolve.",
        "SpringAiEmbeddingAdapter calls Google Text-Embedding-004 to produce a 768-dimensional vector embedding.",
        "PineconeVectorAdapter executes Cosine Similarity Search against the 'servicedesk-knowledge' index (Top-10 candidates).",
        "RerankingEngine uses Cross-Encoder scoring to re-rank chunks and selects the Top-5 most relevant context passages.",
        "SpringAiLlmAdapter prompts Gemini 3.6 Flash with the context passages and formats step-by-step instructions.",
        "ConfidenceCalculator evaluates a multi-factor score (0 - 100) based on similarity distance, source freshness, and keyword overlap.",
        "If score >= threshold (e.g. 75%), return deflectionSuccessful = true with step-by-step solution.",
        "If score < threshold, create ServiceNow incident automatically and return ticket sys_id and INC number."
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("4.2 Multi-Factor Confidence Score Calculator", level=2)
    add_table(doc,
        ["Factor", "Weight", "Description"],
        [
            ["Vector Similarity Score", "50%", "Cosine similarity from Pinecone (0.0 to 1.0)"],
            ["Source Freshness", "20%", "Recency of the referenced knowledge runbook"],
            ["Keyword Exact Overlap", "15%", "Match count of key technical terms (e.g. VPN, OAuth, cert)"],
            ["Source Type Reliability", "15%", "ServiceNow KB (1.0) vs Confluence Page (0.85) vs PDF (0.75)"]
        ]
    )

    # ══════════════════════════════════════════════════════════════
    # 5. REST API SPECIFICATIONS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. REST API Specifications", level=1)

    doc.add_heading("5.1 AI Incident Resolution Endpoint", level=2)
    p = doc.add_paragraph()
    p.add_run("POST ").bold = True
    p.add_run("/api/v1/suggestions/resolve")
    p.add_run("  —  Consumes & Produces: application/json")

    add_table(doc,
        ["Field", "Type", "Required", "Example", "Description"],
        [
            ["title", "String", "Yes", "VPN Connection Failed", "Pre-ticket issue title typed by user"],
            ["description", "String", "Yes", "Server Certificate Invalid error", "Detailed error message or symptoms"],
            ["callerEmail", "String", "No", "user@enterprise.com", "Employee email address"],
            ["userDepartment", "String", "No", "Engineering", "Department context"],
            ["minConfidenceThreshold", "Integer", "No", "75", "Minimum confidence percentage required for deflection"]
        ]
    )

    doc.add_heading("5.2 Knowledge Document Ingestion Endpoint", level=2)
    p = doc.add_paragraph()
    p.add_run("POST ").bold = True
    p.add_run("/api/v1/knowledge/documents")
    p.add_run("  —  Consumes: application/json")

    doc.add_heading("5.3 Analytics & Deflection ROI Endpoint", level=2)
    p = doc.add_paragraph()
    p.add_run("GET ").bold = True
    p.add_run("/api/v1/analytics/deflection")
    p.add_run("  —  Returns current deflection rate %, cost savings USD, and incident metrics.")

    # ══════════════════════════════════════════════════════════════
    # 6. DEPLOYMENT & OPERATIONS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. Deployment & Operations", level=1)

    doc.add_heading("6.1 Containerization Strategy", level=2)
    doc.add_paragraph(
        "The application uses a multi-stage Dockerfile with Maven build optimization and an Alpine Java 21 runtime image. "
        "The image runs under a hardened non-root user (appuser)."
    )

    doc.add_heading("6.2 JVM Performance Tuning for Containers", level=2)
    p = doc.add_paragraph()
    p.add_run("ENV JAVA_OPTS=\"-XX:+UseG1GC -XX:MaxRAMPercentage=75.0 -XX:+ExitOnOutOfMemoryError\"").font.name = "Consolas"

    # ══════════════════════════════════════════════════════════════
    # 7. ROADMAP & POC EXECUTION MILESTONES
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("7. Roadmap & POC Execution Milestones", level=1)

    roadmap = [
        ("Phase 1: Knowledge Ingestion & Vector Indexing", "Set up Apache Tika parsers, SlidingWindowChunker, and Pinecone vector store adapter."),
        ("Phase 2: RAG Pipeline & Gemini 3.6 Synthesis", "Implement Spring AI Gemini adapter, cross-encoder reranking, and multi-factor confidence calculator."),
        ("Phase 3: ServiceNow Integration & Circuit Breaker", "Connect ServiceNow REST API v2, OAuth2 token handling, Resilience4j circuit breaker, and offline queue."),
        ("Phase 4: Deflection Analytics & Observability", "Deploy Prometheus metrics, Grafana deflection dashboard, and AOP audit log tables in PostgreSQL.")
    ]
    for phase, desc in roadmap:
        p = doc.add_paragraph()
        p.add_run(f"{phase}: ").bold = True
        p.add_run(desc)

    # ── Footer ──
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("AI Service Desk Knowledge Intelligence Platform — Confidential  |  ").font.size = Pt(8)
    add_page_number(footer)

    return doc


if __name__ == "__main__":
    output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "output")
    os.makedirs(output_dir, exist_ok=True)

    doc = build_document()
    output_path = os.path.join(output_dir, "AI_Service_Desk_Requirements_Specification.docx")
    doc.save(output_path)
    print(f"Document saved: {output_path}")
