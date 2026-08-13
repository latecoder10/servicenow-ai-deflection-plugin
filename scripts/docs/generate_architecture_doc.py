"""
Generate: AI Service Desk Knowledge Intelligence Platform - Architecture & Technical Specification
POC Document Generator - reflects actual codebase implementation
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os


def setup_page(doc, size="A4"):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(3.18)
    section.orientation = WD_ORIENTATION.PORTRAIT


def tune_styles(doc):
    body = doc.styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(6)
    for n, size, color in [(1, 18, 0x1F3A5F), (2, 14, 0x1F3A5F), (3, 12, 0x2E4057)]:
        s = doc.styles[f"Heading {n}"]
        s.font.name = "Calibri Light"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)
        s.paragraph_format.space_before = Pt(14 - 2 * n)
        s.paragraph_format.space_after = Pt(4)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph("AI Service Desk Knowledge Intelligence Platform", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p = doc.add_paragraph("Architecture & Technical Specification", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph("Version 2.5.0-SNAPSHOT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    p = doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph("CONFIDENTIAL - Proof of Concept")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p = doc.add_paragraph("Prepared by: Estuate Inc.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def add_toc(doc):
    p = doc.add_paragraph("Table of Contents", style="Heading 1")
    p2 = doc.add_paragraph()
    run = p2.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click and choose Update Field to populate."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    for x in (fldChar1, instrText, fldChar2, fldChar3, fldChar4):
        run._r.append(x)


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, name in enumerate(header):
        hdr[i].text = name
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cells[c_idx].text = str(value)
            for p in cells[c_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_endpoint_block(doc, method, path, desc):
    p = doc.add_paragraph()
    p.add_run(f"{method} ").bold = True
    p.add_run(path)
    p.add_run(f"  -  {desc}").italic = True


def build_document():
    doc = Document()
    setup_page(doc)
    tune_styles(doc)
    add_cover(doc)
    doc.add_page_break()
    add_toc(doc)
    doc.add_page_break()

    # ===== 1. PROJECT OVERVIEW =====
    doc.add_heading("1. Project Overview", level=1)

    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "The AI Service Desk Knowledge Intelligence Platform is an enterprise-grade "
        "intelligent incident resolution system that leverages vector-based semantic search, "
        "AI embeddings, and LLM-powered reranking to automatically suggest solutions to IT "
        "support agents when they create new tickets. The platform ingests historical resolved "
        "incidents and knowledge base articles from ServiceNow, processes them through a "
        "multi-stage AI pipeline (chunking, embedding, vector upsert), and provides real-time "
        "AI-powered resolution suggestions through a ServiceNow plugin integration."
    )

    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph(
        "This document covers the complete technical architecture of the AI Service Desk "
        "Knowledge Intelligence Platform, including the Spring Boot backend (Hexagonal "
        "Architecture), the ServiceNow scoped plugin integration, and the optional React "
        "management frontend. It serves as the authoritative reference for the POC demonstration "
        "and future implementation handoff."
    )
    p = doc.add_paragraph()
    p.add_run("Primary Integration Path: ").bold = True
    p.add_run(
        "The end-to-end user interaction flows through the ServiceNow plugin. Users interact "
        "with the Incident form in ServiceNow; the plugin's Client Script triggers a GlideAjax "
        "call to the Script Include (server-side), which makes an outbound REST call to the "
        "Spring Boot backend. There is no direct PDF upload from the end user - all knowledge "
        "ingestion happens through the backend's sync pipeline (ServiceNow records) or the "
        "file upload API (admin-initiated document ingestion)."
    )

    doc.add_heading("1.3 Key Capabilities", level=2)
    capabilities = [
        "Real-time AI-powered incident deflection via ServiceNow plugin (GlideAjax + REST)",
        "Semantic vector search using Google Gemini embeddings (1024-dim) and Pinecone",
        "LLM-powered reranking and resolution generation (Gemini 3.6 Flash)",
        "Multi-stage async document ingestion pipeline (virus scan, storage, chunking, embedding, upsert)",
        "ServiceNow bi-directional integration (incidents, KB articles, attachments) with OAuth2 + Resilience4j",
        "Knowledge connector abstraction with registry pattern for extensibility",
        "PostgreSQL persistence for operational metadata (19 tables via Liquibase)",
        "Confidence scoring with deflection eligibility calculation (60% avg + 30% max + 10% source bonus)",
        "Duplicate detection engine using vector similarity search",
        "Scheduled incremental sync (daily 2 AM cron) with configurable lookback",
        "Docker Compose deployment (PostgreSQL 16, Redis 7, Prometheus, Grafana)",
        "JWT-based security with dev/permissive and prod/stateless profiles",
        "REST API with 20+ endpoints across 7 controllers",
        "Synthetic data loader for POC demonstration",
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style="List Bullet")

    doc.add_heading("1.4 Business Value", level=2)
    for item in [
        "Reduced Mean Time to Resolution (MTTR) through AI-assisted suggestions",
        "Automated incident deflection - agents get resolution suggestions as they type",
        "Knowledge reuse across support teams via vector similarity search",
        "Consistent resolution quality based on historical incident patterns",
        "Lower training costs for new support agents with step-by-step guided resolutions",
        "ROI analytics dashboard for tracking deflection rates and cost savings",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ===== 2. SYSTEM ARCHITECTURE =====
    doc.add_heading("2. System Architecture", level=1)

    doc.add_heading("2.1 High-Level Architecture", level=2)
    doc.add_paragraph(
        "The system follows a Hexagonal Architecture (Ports & Adapters) pattern with clear "
        "separation of concerns across 11 Maven modules. The core domain is isolated from "
        "external integrations via port interfaces, enabling testability and extensibility."
    )
    add_table(doc,
        ["Layer", "Component", "Technology", "Purpose"],
        [
            ["Presentation", "ServiceNow Plugin", "Scoped App (GlideAjax)", "Agent-facing incident form integration"],
            ["Presentation", "React Frontend (Optional)", "React 19 + MUI 9", "Management dashboard and admin UI"],
            ["Gateway", "Tunnel (Dev)", "LocalTunnel / Cloudflare", "Public URL for local backend during dev"],
            ["API", "REST Controllers (7)", "Spring Boot 3.4.2", "20+ REST endpoints across 7 controllers"],
            ["Application", "CQRS Use Cases", "Spring Services", "SuggestionEngine, SyncOrchestrator, Ingestion"],
            ["Domain", "Hexagonal Ports", "Java 17 Interfaces", "EmbeddingPort, VectorDBPort, ServiceNowPort, LlmPort"],
            ["Integration", "Pinecone Adapter", "REST API", "Vector similarity search and storage"],
            ["Integration", "Gemini Embedding Adapter", "Google Gemini API", "Text to 1024-dim vector conversion"],
            ["Integration", "Gemini LLM Adapter", "Gemini 3.6 Flash", "Resolution generation and reranking"],
            ["Integration", "ServiceNow Adapter", "REST + OAuth2", "Bi-directional ServiceNow integration"],
            ["Persistence", "PostgreSQL 16", "JPA + Liquibase", "19 tables for operational metadata"],
            ["Cache", "Redis 7", "Docker", "Distributed cache and rate limiting"],
            ["Observability", "Prometheus + Grafana", "Docker", "Metrics collection and dashboards"],
        ]
    )

    doc.add_heading("2.2 Architecture Diagram", level=2)
    doc.add_paragraph(
        "The following diagram illustrates the complete data flow from ServiceNow user "
        "interaction through the backend AI pipeline to external services:"
    )
    arch_text = """
+-----------------------------------------------------------------------------+
|                    SERVICENow INSTANCE (dev*.service-now.com)               |
|                                                                             |
|   +------------------+    +------------------+    +------------------+     |
|   |  Client Script    |--->|  Script Include   |--->|  REST Message    |     |
|   |  (onChange, 600ms |    |  AIDeflection     |    |  POST to         |     |
|   |   debounce)       |    |  Broker           |    |  Backend API     |     |
|   +------------------+    +------------------+    +--------+---------+     |
|                                                             |               |
|   +------------------+    +------------------+             |               |
|   |  UI Page          |    |  GlideModal       |<------------+               |
|   |  (Jelly XML)      |    |  (500px popup)    |  JSON Response             |
|   +------------------+    +------------------+                            |
+---------------------------------+-------------------------------------------+
                                  |
                                  v
+-----------------------------------------------------------------------------+
|                    LOCALTUNNEL / CLOUDFLARE (Dev Only)                       |
|                    HTTPS Proxy -> localhost:8080                              |
+---------------------------------+-------------------------------------------+
                                  |
                                  v
+-----------------------------------------------------------------------------+
|                    SPRING BOOT BACKEND (localhost:8080)                      |
|                    Hexagonal Architecture - 11 Maven Modules                |
|                                                                             |
|   +----------------+  +----------------+  +----------------+               |
|   |  Suggestion     |  |  Sync           |  |  Ingest         |               |
|   |  Controller     |  |  Controller     |  |  Controller     |               |
|   +-------+--------+  +-------+--------+  +-------+--------+               |
|           |                    |                    |                        |
|   +-------v--------------------v--------------------v--------+              |
|   |              APPLICATION LAYER (CQRS Use Cases)           |              |
|   |  SuggestionEngineService | ServiceNowSyncOrchestrator    |              |
|   |  AsyncDocumentIngestion  | DuplicateDetectionEngine      |              |
|   |  ConfidenceCalculator    | PromptBuilderService           |              |
|   +-------+--------------------+--------------------+--------+              |
|           |                    |                    |                        |
|   +-------v--------+  +-------v--------+  +-------v--------+               |
|   |  Pinecone       |  |  Gemini         |  |  ServiceNow     |               |
|   |  VectorAdapter  |  |  Embedding +    |  |  RestAdapter    |               |
|   |  (Port impl)    |  |  LLM Adapter    |  |  (Resilience4j) |               |
|   +-------+--------+  +-------+--------+  +-------+--------+               |
|           |                    |                    |                        |
|   +-------v--------+  +-------v--------+  +-------v--------+               |
|   |  PostgreSQL     |  |  Liquibase       |  |  Knowledge      |               |
|   |  (19 tables)    |  |  (Schema mgmt)   |  |  Connector      |               |
|   +----------------+  +----------------+  |  Registry       |               |
|                                            +----------------+               |
+----------+----------------------+----------------------+--------------------+
           |                      |                      |
           v                      v                      v
+-----------------+  +-----------------+  +-----------------+
|   Pinecone DB    |  |  Google Gemini   |  |  ServiceNow     |
|   (1024-dim      |  |  Embeddings +    |  |  REST API       |
|    vectors)      |  |  LLM (Flash)     |  |  (OAuth2)       |
+-----------------+  +-----------------+  +-----------------+
"""
    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.name = "Consolas"
    run.font.size = Pt(7)

    doc.add_heading("2.3 Deployment Model", level=2)
    p = doc.add_paragraph()
    p.add_run("Hexagonal Architecture with Docker Compose Infrastructure").bold = True
    doc.add_paragraph(
        "The backend is deployed as a single Spring Boot application with Docker Compose "
        "providing the supporting infrastructure:"
    )
    for r in [
        "Clean separation of domain logic from infrastructure via Ports & Adapters",
        "Externalized state - vectors in Pinecone, embeddings via Gemini API, metadata in PostgreSQL",
        "Observable system - Prometheus metrics with Grafana dashboards",
        "Scalable async processing - thread pool (10-50 threads) for document ingestion",
        "Resilient integrations - Resilience4j circuit breaker + retry for ServiceNow",
        "Cost-effective for POC - no Kubernetes or service mesh required",
    ]:
        doc.add_paragraph(r, style="List Bullet")

    # ===== 3. TECHNOLOGY STACK =====
    doc.add_heading("3. Technology Stack", level=1)

    doc.add_heading("3.1 Core Runtime", level=2)
    add_table(doc,
        ["Component", "Technology", "Version", "Purpose"],
        [
            ["Language", "Java", "17 LTS", "Primary runtime"],
            ["Framework", "Spring Boot", "3.4.2", "Web framework, DI, configuration"],
            ["AI Framework", "Spring AI", "1.0.0-M6", "AI/ML integration abstractions"],
            ["Build Tool", "Apache Maven", "3.9+", "Multi-module dependency management"],
            ["Server", "Embedded Tomcat", "10.1.x", "HTTP server (bundled)"],
            ["ORM", "Hibernate", "JPA", "Object-relational mapping"],
            ["Migrations", "Liquibase", "-", "Database schema management"],
        ]
    )

    doc.add_heading("3.2 AI / ML Integration", level=2)
    add_table(doc,
        ["Component", "Technology", "Version", "Purpose"],
        [
            ["Embedding Model", "Google Gemini", "gemini-embedding-001", "Text to vector conversion"],
            ["Vector Dimension", "1024", "-", "Embedding output size"],
            ["LLM Model", "Google Gemini", "gemini-3.6-flash", "Resolution generation and reranking"],
            ["Batch API", "batchEmbedContents", "v1beta", "Efficient batch processing (100/batch)"],
            ["Reranking", "RerankingEngine", "Internal", "Score-based candidate reranking"],
            ["Confidence", "ConfidenceCalculator", "Internal", "Composite scoring (60+30+10 formula)"],
        ]
    )

    doc.add_heading("3.3 Vector Database", level=2)
    add_table(doc,
        ["Component", "Technology", "Version", "Purpose"],
        [
            ["Vector DB", "Pinecone", "Serverless", "Similarity search and vector storage"],
            ["Index Name", "servicedesk-knowledge", "-", "Vector storage index"],
            ["Dimension", "1024", "-", "Must match embedding model"],
            ["Metric", "cosine", "-", "Similarity measurement"],
            ["Namespace Resolver", "PineconeIndexResolver", "Internal", "5-year time window namespaces"],
            ["Batch Upsert", "96 vectors/call", "-", "Optimized for Pinecone limits"],
        ]
    )

    doc.add_heading("3.4 Persistence & Infrastructure", level=2)
    add_table(doc,
        ["Component", "Technology", "Version", "Purpose"],
        [
            ["Database", "PostgreSQL", "16-alpine", "Operational metadata (19 tables)"],
            ["Cache", "Redis", "7-alpine", "Distributed cache and rate limiting"],
            ["Metrics", "Prometheus", "latest", "Telemetry collection"],
            ["Dashboards", "Grafana", "latest", "Analytics visualization"],
            ["Containerization", "Docker Compose", "3.8", "Infrastructure orchestration"],
        ]
    )

    doc.add_heading("3.5 ServiceNow Integration", level=2)
    add_table(doc,
        ["Component", "Technology", "Version", "Purpose"],
        [
            ["Instance", "ServiceNow", "Xanadu", "ITSM platform"],
            ["Auth", "OAuth2", "Bearer Token", "API authentication (password or client_credentials)"],
            ["Plugin", "Scoped Application", "x_2185757_ai_tic_0", "AI Ticket Deflection app"],
            ["Plugin", "Script Include", "-", "Server-side GlideAjax broker (AIDeflectionBroker)"],
            ["Plugin", "Client Script", "-", "Browser-side onChange (600ms debounce)"],
            ["Plugin", "REST Message", "-", "Outbound HTTP to Spring Boot backend"],
            ["Plugin", "UI Page", "Jelly XML", "Modal popup for resolution display"],
            ["Resilience", "Resilience4j", "2.2.0", "Circuit breaker + retry for ServiceNow API"],
        ]
    )

    doc.add_heading("3.6 DevOps & Security", level=2)
    add_table(doc,
        ["Component", "Technology", "Version", "Purpose"],
        [
            ["Tunneling", "LocalTunnel / Cloudflare", "-", "Public URL for local backend (dev only)"],
            ["API Docs", "SpringDoc OpenAPI", "2.8.4", "Swagger UI at /swagger-ui.html"],
            ["Security", "Spring Security", "-", "JWT auth (prod) / permissive (dev)"],
            ["JWT Library", "JJWT", "0.12.6", "Token generation and validation"],
            ["Resilience", "Resilience4j", "2.2.0", "Circuit breaker, retry, rate limiting"],
            ["Object Mapping", "MapStruct", "1.6.3", "Entity-to-DTO mapping"],
            ["Code Gen", "Lombok", "1.18.36", "Boilerplate reduction"],
            ["Doc Parsing", "Apache Tika", "3.1.0", "Text extraction from PDF, DOCX, etc."],
        ]
    )

    # ===== 4. MODULE STRUCTURE =====
    doc.add_heading("4. Module Structure", level=1)

    doc.add_heading("4.1 Multi-Module Maven Project", level=2)
    doc.add_paragraph(
        "The project follows a clean multi-module Maven reactor with 11 sub-modules "
        "organized in a Hexagonal Architecture pattern. The `api` module depends on all "
        "other modules and serves as the Spring Boot entry point."
    )
    add_table(doc,
        ["Module", "Artifact ID", "Package", "Purpose"],
        [
            ["common", "common", "com.servicedesk.ai.common", "Base entities, exceptions, correlation context, ProblemDetails"],
            ["domain", "domain", "com.servicedesk.ai.domain", "Core entities, value objects, port interfaces, AppConstants"],
            ["application", "application", "com.servicedesk.ai.application", "CQRS use cases, services, connectors, orchestrators"],
            ["knowledge-loader", "knowledge-loader", "com.servicedesk.ai.loader", "Document parsing (Tika), chunking (sliding window), file storage"],
            ["pinecone", "pinecone-integration", "com.servicedesk.ai.integration.pinecone", "Pinecone vector DB adapter, index namespace resolver"],
            ["servicenow", "servicenow-integration", "com.servicedesk.ai.integration.servicenow", "ServiceNow REST adapter with OAuth2 + Resilience4j"],
            ["llm", "llm-integration", "com.servicedesk.ai.integration.llm", "Gemini embedding adapter, LLM adapter, reranking engine"],
            ["analytics", "analytics", "com.servicedesk.ai.analytics", "Deflection analytics and ROI metric computation"],
            ["security", "security", "com.servicedesk.ai.security", "JWT token provider, auth filter, Spring Security config"],
            ["infrastructure", "infrastructure", "com.servicedesk.ai.infrastructure", "ServiceNow sync scheduler, AOP audit logging aspect"],
            ["api", "api", "com.servicedesk.ai.api", "REST controllers, DTOs, exception handler, Spring Boot main class"],
        ]
    )

    doc.add_heading("4.2 Dependency Flow", level=2)
    doc.add_paragraph(
        "Dependencies flow inward following Hexagonal Architecture principles:"
    )
    for item in [
        "common <- domain (base classes, exceptions)",
        "domain <- application (use cases reference domain interfaces)",
        "domain <- knowledge-loader (chunking operates on domain types)",
        "domain <- pinecone/servicenow/llm (adapters implement port interfaces)",
        "domain <- analytics (metrics computed from domain models)",
        "common <- security (JWT uses base exceptions)",
        "domain, application <- infrastructure (schedulers, aspects)",
        "ALL modules <- api (wires up all adapters, entry point)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4.3 Domain Ports (Hexagonal Interfaces)", level=2)
    doc.add_paragraph(
        "The domain module defines 6 output port interfaces that integration adapters must implement:"
    )
    add_table(doc,
        ["Port Interface", "Methods", "Implementing Adapter"],
        [
            ["EmbeddingPort", "generateEmbedding(), generateBatchEmbeddings()", "SpringAiEmbeddingAdapter"],
            ["VectorDatabasePort", "upsertChunks(), similaritySearch(), deleteByDocumentId(), countVectors()", "PineconeVectorAdapter"],
            ["ServiceNowPort", "createIncident(), fetchResolvedIncidentsSince(), downloadAttachmentContent(), validateConnection()", "ServiceNowRestAdapter"],
            ["LlmPort", "generateResolution(), rerank()", "SpringAiLlmAdapter"],
            ["KnowledgeConnector", "getConnectorType(), testConnection(), synchronize(), fetchChanges()", "ServiceNowKnowledgeConnector"],
            ["KnowledgeRepositoryPort", "saveDocumentMetadata(), findDocumentById(), saveFeedback()", "KnowledgeDocumentJpaRepository"],
        ]
    )

    # ===== 5. AI/ML PIPELINE =====
    doc.add_heading("5. AI/ML Pipeline", level=1)

    doc.add_heading("5.1 Knowledge Ingestion Pipeline", level=2)
    doc.add_paragraph(
        "The ingestion pipeline processes knowledge from multiple sources through a "
        "multi-stage async pipeline. There are two ingestion paths:"
    )
    p = doc.add_paragraph()
    p.add_run("Path A - ServiceNow Sync (Primary): ").bold = True
    p.add_run("Scheduled daily at 2 AM or manual trigger via API")

    add_table(doc,
        ["Step", "Component", "Action", "Output"],
        [
            ["1", "ServiceNowSyncScheduler", "Triggers sync (cron or manual)", "Sync request"],
            ["2", "ServiceNowKnowledgeConnector", "Fetches resolved incidents + KB articles via REST", "KnowledgeRecord list"],
            ["3", "TextChunker", "Splits text into 1500-char chunks (200-char overlap)", "Text chunks"],
            ["4", "SpringAiEmbeddingAdapter", "Generates 1024-dim vectors via Gemini batch API (100/batch)", "Vector embeddings"],
            ["5", "PineconeVectorAdapter", "Upserts vectors in batches of 96 with exponential backoff", "Stored in Pinecone"],
            ["6", "SyncJobJpaRepository", "Persists job status and metrics to PostgreSQL", "Audit trail"],
        ]
    )

    p = doc.add_paragraph()
    p.add_run("Path B - Document Upload (Admin-Initiated): ").bold = True
    p.add_run("Files uploaded via REST API go through 6-stage async ingestion")

    add_table(doc,
        ["Step", "Component", "Action", "Output"],
        [
            ["1", "AsyncDocumentIngestionService", "Virus scan (NoOp in POC)", "Scan result"],
            ["2", "LocalFileStorageService", "Stores file to local filesystem", "Storage path"],
            ["3", "KnowledgeDocumentJpaRepository", "Persists document metadata to PostgreSQL", "DocumentEntity"],
            ["4", "TextDocumentParser (Tika)", "Extracts text from PDF, DOCX, XLSX, CSV, TXT, MD", "Raw text"],
            ["5", "SlidingWindowChunker", "Chunks text (300 words, 50 overlap)", "Text chunks"],
            ["6", "PineconeVectorAdapter", "Embeds and upserts to Pinecone", "Vector index"],
        ]
    )

    doc.add_heading("5.2 Embedding Generation", level=2)
    doc.add_paragraph(
        "The system uses Google Gemini's gemini-embedding-001 model to convert text into "
        "1024-dimensional vectors for semantic search:"
    )
    add_table(doc,
        ["Parameter", "Value", "Purpose"],
        [
            ["Model", "gemini-embedding-001", "Google's embedding model"],
            ["Dimension", "1024", "Vector output size"],
            ["Batch Size", "100 texts per call", "Optimized for Gemini API limits"],
            ["Max Text Length", "20,000 characters", "Truncation limit"],
            ["Fallback", "Hash-based placeholder", "When API key is missing (dev mode)"],
        ]
    )

    doc.add_heading("5.3 Similarity Search & Reranking", level=2)
    doc.add_paragraph(
        "When a query arrives (from ServiceNow plugin or API), the system performs a "
        "multi-stage retrieval and ranking process:"
    )
    for step in [
        "1. Agent types in Short description field (ServiceNow) or submits via API",
        "2. Client Script debounces input (600ms) and fires GlideAjax to Script Include",
        "3. Script Include makes outbound REST POST to /api/v1/suggestions/resolve",
        "4. SuggestionEngineService generates query embedding via Gemini",
        "5. Pinecone performs cosine similarity search (top-K results)",
        "6. RerankingEngine sorts candidates by relevance score",
        "7. ConfidenceCalculator computes composite score: 60% avg + 30% max + 10% source count",
        "8. PromptBuilderService constructs LLM prompt with retrieved knowledge context",
        "9. SpringAiLlmAdapter generates resolution via Gemini 3.6 Flash",
        "10. Response returned as SuggestionResponse JSON to ServiceNow plugin",
        "11. Script Include passes JSON to Client Script via GlideAjax callback",
        "12. Client Script opens GlideModal with UI Page showing resolution popup",
    ]:
        doc.add_paragraph(step, style="List Bullet")

    doc.add_heading("5.4 Confidence Scoring", level=2)
    add_table(doc,
        ["Score Range", "Band", "Interpretation", "Deflection Eligible"],
        [
            ["85 - 100", "VERY_HIGH", "Excellent match", "Yes"],
            ["70 - 84", "HIGH", "Good match", "Yes"],
            ["50 - 69", "MEDIUM", "Moderate match", "No"],
            ["0 - 49", "LOW", "Low match", "No"],
        ]
    )

    doc.add_heading("5.5 Duplicate Detection", level=2)
    doc.add_paragraph(
        "The DuplicateDetectionEngine uses vector similarity search against Pinecone to "
        "detect similar resolved incidents before a new ticket is created, preventing "
        "duplicate knowledge entries."
    )

    # ===== 6. SERVICENOW PLUGIN ARCHITECTURE =====
    doc.add_heading("6. ServiceNow Plugin Architecture", level=1)

    doc.add_heading("6.1 Scoped Application", level=2)
    doc.add_paragraph(
        "The ServiceNow integration is packaged as a Scoped Application named "
        "\"AI Ticket Deflection\" within the x_2185757_ai_tic_0 scope."
    )
    add_table(doc,
        ["Component", "Record Name", "Type", "Purpose"],
        [
            ["Script Include", "AIDeflectionBroker", "Server-side JS", "GlideAjax handler - bridges browser to Spring Boot REST API"],
            ["Client Script", "AI Ticket Deflection Listener", "onChange", "Detects Short description changes, triggers AI search"],
            ["REST Message", "Spring Boot Deflection API", "Outbound REST", "HTTP POST to Spring Boot /api/v1/suggestions/resolve"],
            ["UI Page", "ai_resolution_popup", "Jelly XML", "Modal popup displaying AI resolution suggestion"],
        ]
    )

    doc.add_heading("6.2 Script Include - AIDeflectionBroker (Server-Side)", level=2)
    doc.add_paragraph(
        "The Script Include extends AbstractAjaxProcessor and is the ONLY component "
        "that makes outbound HTTP calls. The browser never calls Spring Boot directly."
    )
    add_table(doc,
        ["Function", "Parameters", "Returns", "Purpose"],
        [
            ["getResolution", "sysparm_title, sysparm_description", "JSON string", "Fetches AI resolution from Spring Boot backend"],
        ]
    )
    for d in [
        "Extends AbstractAjaxProcessor for GlideAjax compatibility",
        "Uses sn_ws.RESTMessageV2 referencing the scoped REST Message record",
        "Sets template variables: title, description, callerEmail (from gs.getUser().getEmail())",
        "Returns raw JSON response from Spring Boot to Client Script",
        "Catches exceptions and returns error JSON instead of crashing",
    ]:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("6.3 Client Script - AI Ticket Deflection Listener (Browser-Side)", level=2)
    doc.add_paragraph(
        "The Client Script runs in the browser and triggers on Short description field "
        "changes in the Incident form. It uses a 600ms debounce to prevent excessive "
        "backend calls during typing."
    )
    add_table(doc,
        ["Function", "Trigger", "Behavior"],
        [
            ["onChange", "Short description changes", "Guards on loading/empty, starts 600ms debounce timer"],
            ["GlideAjax callback", "After debounce", "Creates GlideAjax to AIDeflectionBroker.getResolution"],
            ["Response parsing", "On response", "Parses JSON, extracts summary, steps, code snippet"],
            ["GlideModal", "On valid response", "Opens 500px modal with UI Page (ai_resolution_popup)"],
        ]
    )

    doc.add_heading("6.4 UI Page - ai_resolution_popup (Modal Content)", level=2)
    doc.add_paragraph(
        "The UI Page renders inside a GlideModal popup and displays the AI resolution "
        "suggestion using ServiceNow's Jelly XML templating engine."
    )
    for feature in [
        "Resolution Summary - teal-bordered highlight box with AI-generated summary",
        "Recommended Steps - numbered troubleshooting steps",
        "Terminal Command / Snippet - dark-themed code block (conditionally shown)",
        "Data passed via setPreference() / RP.getWindowProperties()",
        "Modal is draggable and non-blocking - Incident form remains fully usable",
    ]:
        doc.add_paragraph(feature, style="List Bullet")

    doc.add_heading("6.5 End-to-End Plugin Data Flow", level=2)
    plugin_flow = """
User types in Short description field
        |
        v
Client Script (onChange, 600ms debounce)
        |
        v
GlideAjax -> Script Include (AIDeflectionBroker.getResolution)
        |
        v
RESTMessageV2 -> POST /api/v1/suggestions/resolve (via LocalTunnel)
        |
        v
Spring Boot Backend (SuggestionController -> SuggestionEngineService)
        |
        v
Gemini Embedding -> Pinecone Similarity Search -> Reranking -> LLM Resolution
        |
        v
JSON Response -> Script Include -> GlideAjax Callback
        |
        v
GlideModal -> UI Page (ai_resolution_popup) -> User sees resolution
"""
    p = doc.add_paragraph()
    run = p.add_run(plugin_flow)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # ===== 7. DATA FLOW =====
    doc.add_heading("7. Data Flow", level=1)

    doc.add_heading("7.1 Ingestion Flow (Historical Data Sync)", level=2)
    add_table(doc,
        ["Step", "Component", "Action", "Output"],
        [
            ["1", "ServiceNowSyncScheduler", "Triggers daily at 2 AM", "Sync request"],
            ["2", "ServiceNowKnowledgeConnector", "Fetches resolved incidents + KB articles", "KnowledgeRecord list"],
            ["3", "TextChunker", "Splits text into 1500-char chunks (200-char overlap)", "Text chunks"],
            ["4", "SpringAiEmbeddingAdapter", "Generates 1024-dim vectors via Gemini batch API", "Vector embeddings"],
            ["5", "PineconeVectorAdapter", "Upserts vectors in batches of 96 with retry", "Stored vectors"],
            ["6", "SyncJobJpaRepository", "Persists job metrics to PostgreSQL", "Audit trail"],
        ]
    )

    doc.add_heading("7.2 Query Flow (Real-Time Suggestions via Plugin)", level=2)
    add_table(doc,
        ["Step", "Component", "Action", "Output"],
        [
            ["1", "Client Script", "Detects Short description change (600ms debounce)", "Debounced input"],
            ["2", "GlideAjax", "Calls AIDeflectionBroker.getResolution", "Server request"],
            ["3", "Script Include", "POSTs to backend via RESTMessageV2", "REST request"],
            ["4", "SuggestionEngineService", "Orchestrates full AI pipeline", "Query processing"],
            ["5", "SpringAiEmbeddingAdapter", "Embeds query text via Gemini", "Query vector"],
            ["6", "PineconeVectorAdapter", "Performs cosine similarity search (top-K)", "Matching vectors"],
            ["7", "RerankingEngine", "Sorts candidates by relevance score", "Reranked results"],
            ["8", "ConfidenceCalculator", "Computes composite confidence score", "Confidence band"],
            ["9", "PromptBuilderService", "Builds LLM prompt with retrieved context", "LLM prompt"],
            ["10", "SpringAiLlmAdapter", "Generates resolution via Gemini 3.6 Flash", "AI resolution"],
            ["11", "Response", "Returns SuggestionResponse JSON", "JSON response"],
            ["12", "Client Script", "Opens GlideModal with UI Page", "UI update"],
        ]
    )

    doc.add_heading("7.3 Document Upload Flow (Admin-Initiated)", level=2)
    add_table(doc,
        ["Step", "Component", "Action", "Output"],
        [
            ["1", "FileStorageController", "Receives multipart file upload", "UploadJobEntity"],
            ["2", "AsyncDocumentIngestionService", "6-stage async pipeline", "Background processing"],
            ["3", "NoOpVirusScanService", "Virus scan (no-op in POC)", "Scan passed"],
            ["4", "LocalFileStorageService", "Stores file to ./storage/documents/", "File path"],
            ["5", "TextDocumentParser (Tika)", "Extracts text from documents", "Raw text"],
            ["6", "SlidingWindowChunker", "Chunks text (300 words, 50 overlap)", "Text chunks"],
            ["7", "SpringAiEmbeddingAdapter", "Generates embeddings via Gemini", "1024-dim vectors"],
            ["8", "PineconeVectorAdapter", "Upserts vectors to Pinecone index", "Stored vectors"],
        ]
    )

    # ===== 8. API SPECIFICATION =====
    doc.add_heading("8. API Specification", level=1)
    doc.add_paragraph(
        "The backend exposes 20+ REST endpoints across 7 controllers. All endpoints "
        "are prefixed with /api/v1. Full OpenAPI docs at /swagger-ui.html."
    )

    doc.add_heading("8.1 Suggestion Controller (/api/v1/suggestions)", level=2)
    add_endpoint_block(doc, "POST", "/api/v1/suggestions/resolve", "Get AI-powered resolution suggestion")
    add_table(doc,
        ["Parameter", "Type", "Required", "Description"],
        [
            ["title", "string", "Yes", "Incident short description"],
            ["description", "string", "Yes", "Incident detailed description"],
            ["callerEmail", "string", "No", "Requesting user's email"],
            ["userDepartment", "string", "No", "User's department"],
            ["category", "string", "No", "Incident category"],
            ["minConfidenceThreshold", "number", "No", "Minimum confidence score (0-100)"],
        ]
    )

    doc.add_heading("8.2 ServiceNow Controller (/api/v1/servicenow)", level=2)
    for method, path, desc in [
        ("POST", "/api/v1/servicenow/incidents", "Create a ServiceNow incident"),
        ("GET", "/api/v1/servicenow/health", "Validate ServiceNow connection health"),
        ("POST", "/api/v1/servicenow/sync/incremental", "Trigger incremental sync pipeline"),
        ("GET", "/api/v1/servicenow/attachments/metadata/{attachmentId}", "Fetch attachment metadata"),
        ("GET", "/api/v1/servicenow/attachments/download/{attachmentId}", "Download attachment binary"),
    ]:
        add_endpoint_block(doc, method, path, desc)

    doc.add_heading("8.3 Knowledge Controller (/api/v1/knowledge)", level=2)
    for method, path, desc in [
        ("GET", "/api/v1/knowledge/search", "Semantic search across Pinecone knowledge index"),
        ("GET", "/api/v1/knowledge/records", "List synchronized knowledge records"),
        ("POST", "/api/v1/knowledge/records/{sysId}/reindex", "Re-embed and upsert a single record"),
        ("DELETE", "/api/v1/knowledge/records/{sysId}", "Delete a vector chunk from Pinecone"),
        ("POST", "/api/v1/knowledge/load-synthetic", "Load 20 sample incidents into ServiceNow"),
    ]:
        add_endpoint_block(doc, method, path, desc)

    doc.add_heading("8.4 File Storage Controller (/api/v1/files)", level=2)
    for method, path, desc in [
        ("POST", "/api/v1/files/upload", "Upload file for async ingestion pipeline"),
        ("GET", "/api/v1/files/jobs/{jobId}", "Get async ingestion progress"),
        ("GET", "/api/v1/files", "List all stored knowledge documents"),
        ("GET", "/api/v1/files/{documentId}/download", "Download stored file"),
    ]:
        add_endpoint_block(doc, method, path, desc)

    doc.add_heading("8.5 Connector Controller (/api/v1/connectors)", level=2)
    for method, path, desc in [
        ("GET", "/api/v1/connectors", "List all registered connector types"),
        ("POST", "/api/v1/connectors/{type}/test", "Test connection health"),
        ("POST", "/api/v1/connectors/{type}/sync", "Trigger async sync job"),
        ("GET", "/api/v1/connectors/{type}/history", "View sync job history"),
    ]:
        add_endpoint_block(doc, method, path, desc)

    doc.add_heading("8.6 Analytics & Pipeline Controllers", level=2)
    for method, path, desc in [
        ("GET", "/api/v1/analytics/deflection", "Get ROI, deflection rate, and metrics"),
        ("GET", "/api/v1/analytics/dashboard", "Full executive dashboard telemetry"),
        ("GET", "/api/v1/pipeline/jobs", "List recent sync pipeline execution jobs"),
        ("GET", "/api/v1/pipeline/jobs/{jobId}", "Get detailed status for a specific job"),
        ("GET", "/api/v1/health", "Platform health check"),
    ]:
        add_endpoint_block(doc, method, path, desc)

    # ===== 9. DATABASE SCHEMA =====
    doc.add_heading("9. Database Schema", level=1)
    doc.add_paragraph(
        "PostgreSQL 16 stores all operational metadata. Schema managed by Liquibase with 19 tables:"
    )
    add_table(doc,
        ["Table", "Purpose", "Key Columns"],
        [
            ["knowledge_documents", "Uploaded knowledge document metadata", "workspaceId, title, sourceType, storagePath, status"],
            ["knowledge_chunks", "Text chunks with embeddings", "documentId, chunkIndex, chunkText, tokenCount"],
            ["sync_jobs", "ServiceNow sync execution history", "jobId, connectorType, status, itemsFetched/Created/Failed"],
            ["upload_jobs", "File upload async processing", "documentId, status, progressPercentage"],
            ["indexing_jobs", "Document indexing progress", "documentId, chunksProcessed, totalChunks"],
            ["connector_configurations", "Data source configs (encrypted secrets)", "connectorType, instanceUrl, isActive"],
            ["audit_logs", "System audit trail (JSONB details)", "eventType, principal, action, resourceId"],
            ["attachment_metadata", "ServiceNow attachment references", "attachmentId, fileName, tableName, recordSysId"],
            ["workspaces", "Organizational workspaces", "name, description"],
            ["departments", "Department hierarchy", "workspaceId, code, name"],
            ["knowledge_categories", "Knowledge classification", "workspaceId, name, slug"],
        ]
    )

    # ===== 10. CONFIGURATION =====
    doc.add_heading("10. Configuration", level=1)

    doc.add_heading("10.1 Environment Variables", level=2)
    add_table(doc,
        ["Variable", "Default", "Description"],
        [
            ["GEMINI_API_KEY", "-", "Google Gemini API key (embeddings + LLM)"],
            ["AI_PINECONE_API_KEY", "-", "Pinecone API key"],
            ["AI_PINECONE_HOST", "-", "Pinecone index host URL"],
            ["SERVICENOW_INSTANCE_URL", "-", "ServiceNow instance URL"],
            ["SERVICENOW_CLIENT_ID", "-", "ServiceNow OAuth2 client ID"],
            ["SERVICENOW_CLIENT_SECRET", "-", "ServiceNow OAuth2 client secret"],
            ["SERVICENOW_USERNAME", "-", "ServiceNow admin username"],
            ["SERVICENOW_PASSWORD", "-", "ServiceNow admin password"],
            ["SPRING_DATASOURCE_URL", "jdbc:postgresql://localhost:5432/servicedesk_ai", "PostgreSQL connection"],
            ["SPRING_DATASOURCE_USERNAME", "postgres", "Database user"],
            ["SPRING_DATASOURCE_PASSWORD", "changeme", "Database password"],
            ["SCHEDULER_SERVICENOW_ENABLED", "true", "Enable daily sync scheduler"],
            ["SCHEDULER_SERVICENOW_CRON", "0 0 2 * * ?", "Sync schedule (daily 2 AM)"],
            ["STORAGE_ROOT_PATH", "./storage", "Local file storage root"],
            ["SERVER_PORT", "8080", "Backend server port"],
        ]
    )

    doc.add_heading("10.2 AppConstants.java", level=2)
    add_table(doc,
        ["Category", "Constant", "Value"],
        [
            ["Vector IDs", "VECTOR_ID_PREFIX", "sn-"],
            ["Embedding", "EMBEDDING_MODEL", "gemini-embedding-001"],
            ["Embedding", "EMBEDDING_DIMENSION", "1024"],
            ["Embedding", "EMBEDDING_BATCH_SIZE", "100"],
            ["Pinecone", "PINECONE_BATCH_SIZE", "96"],
            ["Pinecone", "PINECONE_MAX_RETRIES", "3"],
            ["Chunking", "CHUNK_SIZE_CHARS", "1500"],
            ["Chunking", "CHUNK_OVERLAP_CHARS", "200"],
            ["Defaults", "DEFAULT_WORKSPACE", "Enterprise IT"],
            ["Defaults", "DEFAULT_DEPARTMENT", "Global Service Desk"],
            ["Defaults", "DEFAULT_PRIORITY", "3 - Moderate"],
            ["Connector", "CONNECTOR_SERVICENOW", "SERVICENOW"],
        ]
    )

    doc.add_heading("10.3 Resilience4j Configuration", level=2)
    add_table(doc,
        ["Pattern", "Setting", "Value"],
        [
            ["Circuit Breaker", "slidingWindowSize", "10"],
            ["Circuit Breaker", "failureRateThreshold", "50%"],
            ["Circuit Breaker", "waitDurationInOpenState", "30s"],
            ["Circuit Breaker", "permittedInHalfOpen", "3"],
            ["Retry", "maxAttempts", "3"],
            ["Retry", "waitDuration", "1s (exponential backoff)"],
        ]
    )

    # ===== 11. FRONTEND (ADDITIONAL WORK) =====
    doc.add_heading("11. Frontend - Management Dashboard (Additional Work)", level=1)

    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run(
        "The frontend was developed as additional work beyond the original ServiceNow plugin "
        "integration scope. It provides a management dashboard for monitoring and administering "
        "the AI Service Desk platform. The primary end-user interaction remains through the "
        "ServiceNow plugin."
    )

    doc.add_heading("11.1 Technology Stack", level=2)
    add_table(doc,
        ["Component", "Technology", "Version"],
        [
            ["UI Framework", "React", "19.0.1"],
            ["Language", "TypeScript", "~5.8.2"],
            ["Build Tool", "Vite", "6.2.3"],
            ["UI Library", "MUI (Material UI)", "9.3.1"],
            ["Styling", "Tailwind CSS + Emotion", "4.1.14"],
            ["State Management", "Zustand", "5.0.14"],
            ["HTTP Client", "Axios", "1.19.0"],
            ["Charts", "Recharts", "3.10.1"],
            ["Routing", "react-router-dom", "7.18.2"],
        ]
    )

    doc.add_heading("11.2 Pages and Features", level=2)
    add_table(doc,
        ["Route", "Page", "Purpose"],
        [
            ["/dashboard", "DashboardPage", "Executive dashboard with 4 tabs (Overview, AI Analytics, System Telemetry, Pipeline)"],
            ["/knowledge", "KnowledgePage", "Pinecone knowledge base - search, reindex, delete records"],
            ["/suggestions", "SuggestionsPage", "AI ticket deflection engine - submit queries, get AI resolutions"],
            ["/connectors", "ConnectorsPage", "Enterprise data connectors - test, sync, view history"],
            ["/pipeline", "PipelinePage", "Vector ingestion pipeline - monitor sync jobs, chunking, embedding"],
            ["/files", "FilesPage", "Knowledge document upload - drag-and-drop, progress polling"],
            ["/settings", "SettingsPage", "System settings - platform health, ServiceNow, AI engine status"],
        ]
    )

    doc.add_heading("11.3 Architecture Highlights", level=2)
    for f in [
        "Feature-based directory structure (features/* with co-located components/ and hooks/)",
        "Typed API layer with discriminated union error handling (ApiResult<T>)",
        "Zustand stores for app state, auth (JWT), and notifications",
        "18 reusable UI components in components/ui/",
        "Custom hooks: useApiCall, useDebounce, usePagination, useToast",
        "Vite dev proxy to localhost:8080 for seamless backend integration",
        "Dark mode support via MUI theme configuration",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    # ===== 12. DEPLOYMENT =====
    doc.add_heading("12. Deployment & Operations", level=1)

    doc.add_heading("12.1 Prerequisites", level=2)
    for p_item in [
        "Java 17+ runtime (Eclipse Temurin recommended)",
        "Maven 3.9+ for building from source",
        "Docker & Docker Compose for infrastructure",
        "Node.js 18+ for LocalTunnel (dev) or Cloudflare tunnel",
        "ServiceNow developer instance with admin access",
        "Google Cloud project with Gemini API enabled",
        "Pinecone account (free tier OK)",
    ]:
        doc.add_paragraph(p_item, style="List Bullet")

    doc.add_heading("12.2 Docker Compose Infrastructure", level=2)
    add_table(doc,
        ["Service", "Image", "Port", "Purpose"],
        [
            ["service-desk-ai-api", "Custom (Dockerfile)", "8080, 8081", "Spring Boot backend"],
            ["postgres", "postgres:16-alpine", "5432", "Relational database"],
            ["redis", "redis:7-alpine", "6379", "Distributed cache"],
            ["prometheus", "prom/prometheus:latest", "9090", "Metrics collection"],
            ["grafana", "grafana/grafana:latest", "3001", "Analytics dashboards"],
        ]
    )

    doc.add_heading("12.3 Deployment Steps", level=2)
    add_table(doc,
        ["Step", "Command", "Purpose"],
        [
            ["1", "cd service-desk-ai-platform-backend", "Navigate to backend directory"],
            ["2", "mvn clean install -DskipTests", "Build all 11 modules"],
            ["3", "docker-compose up -d", "Start PostgreSQL, Redis, Prometheus, Grafana"],
            ["4", "mvn spring-boot:run -pl modules/api", "Start the backend application"],
            ["5", "cloudflared tunnel --url http://localhost:8080", "Start tunnel for ServiceNow plugin"],
            ["6", "curl http://localhost:8080/api/v1/health", "Verify backend is running"],
            ["7", "POST /api/v1/knowledge/load-synthetic", "Load POC sample data"],
            ["8", "POST /api/v1/knowledge/records/{id}/reindex", "Embed and index records to Pinecone"],
        ]
    )

    doc.add_heading("12.4 ServiceNow Plugin Setup", level=2)
    add_table(doc,
        ["Step", "Navigation", "Key Settings"],
        [
            ["1", "ServiceNow Studio > Create Application", "Name: AI Ticket Deflection, Scope: x_2185757_ai_tic_0"],
            ["2", "System Web Services > Outbound > REST Message", "Name: Spring Boot Deflection API, Base URL: tunnel URL"],
            ["3", "REST Message > HTTP Method", "Name: resolve, Method: POST, Endpoint: /api/v1/suggestions/resolve"],
            ["4", "System Definition > Script Includes", "Name: AIDeflectionBroker, Client Callable: checked"],
            ["5", "System Definition > Client Scripts", "Type: onChange, Field: Short description"],
            ["6", "Service Portal > UI Pages", "Name: ai_resolution_popup, Type: Jelly XML"],
        ]
    )

    # ===== 13. TESTING =====
    doc.add_heading("13. Testing Strategy", level=1)

    doc.add_heading("13.1 Test Levels", level=2)
    add_table(doc,
        ["Level", "Scope", "Tools"],
        [
            ["Unit", "Individual components (services, adapters)", "JUnit 5, Mockito"],
            ["Integration", "API endpoints, database operations", "MockMvc, TestRestTemplate, Testcontainers"],
            ["E2E", "Full user flow via ServiceNow plugin", "Manual testing in ServiceNow instance"],
        ]
    )

    doc.add_heading("13.2 Verification Checklist", level=2)
    for item in [
        "Backend health check returns UP",
        "PostgreSQL connection established (Liquibase migrations applied)",
        "Pinecone connection validated",
        "Gemini API key configured and embeddings generating",
        "ServiceNow OAuth2 authentication working",
        "Cloudflare/LocalTunnel is accessible",
        "ServiceNow Script Include is client-callable",
        "Client Script triggers on Short description change (600ms debounce)",
        "GlideAjax calls return 200 status",
        "AI resolution modal appears in Incident form",
        "Synthetic data loads successfully into ServiceNow",
        "Sync pipeline indexes records to Pinecone",
        "Frontend dashboard displays at http://localhost:3000",
        "Grafana dashboards show metrics at http://localhost:3001",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ===== 14. TROUBLESHOOTING =====
    doc.add_heading("14. Troubleshooting", level=1)

    doc.add_heading("14.1 Common Issues", level=2)
    add_table(doc,
        ["Issue", "Symptoms", "Fix"],
        [
            ["Backend not running", "Health check fails, tunnel returns 502", "Run: mvn spring-boot:run -pl modules/api"],
            ["Tunnel URL changed", "ServiceNow returns 403/500", "Update REST Message endpoint in ServiceNow"],
            ["Script Include not callable", "GlideAjax fails silently", "Check Client Callable checkbox, verify scope"],
            ["Scoped REST Message name wrong", "Error constructing REST Message", "Use exact scoped name: x_2185757_ai_tic_0.Spring Boot Deflection API"],
            ["No suggestions appear", "Empty modal", "Load synthetic data and reindex to Pinecone"],
            ["Pinecone connection failed", "Vector search error", "Check AI_PINECONE_API_KEY and AI_PINECONE_HOST"],
            ["Gemini API error", "Embedding failure", "Check GEMINI_API_KEY"],
            ["PostgreSQL connection refused", "Liquibase migration fails", "Run: docker-compose up -d postgres"],
            ["Scoped timer issue (window)", "Debounce not working", "Use control.aiDeflectionTimer (not window)"],
        ]
    )

    doc.add_heading("14.2 Debug Steps", level=2)
    for step in [
        "Check browser Console (F12) for JavaScript errors",
        "Check ServiceNow System Logs for [AI Deflection] errors",
        "Verify tunnel URL is correct and accessible",
        "Test API endpoint directly via curl or Postman",
        "Check backend logs for connection errors",
        "Verify REST Message name matches scoped application scope",
        "Check Docker containers are running: docker-compose ps",
        "Verify PostgreSQL is accessible on port 5432",
    ]:
        doc.add_paragraph(step, style="List Bullet")

    # ===== 15. KNOWN LIMITATIONS =====
    doc.add_heading("15. Known Limitations & POC Notes", level=1)

    doc.add_heading("15.1 Current Limitations", level=2)
    for item in [
        "LLM adapter uses hardcoded demo responses (keyword matching) instead of live Gemini LLM calls in dev mode",
        "Virus scan is a NoOp implementation (not production-ready)",
        "UI Page rendering via Jelly RP.getWindowProperties() was not conclusively verified in all scenarios",
        "REST Message template defines userDepartment, category, minConfidenceThreshold but Script Include only populates title, description, callerEmail",
        "LocalTunnel is for development only - production requires proper HTTPS backend URL with OAuth2/mTLS",
        "Frontend auth store is wired for JWT but no login UI page exists",
        "Redis is provisioned in Docker Compose but not actively used in current POC code",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("15.2 POC Demo Flow", level=2)
    doc.add_paragraph(
        "For the POC demonstration, the recommended flow is:"
    )
    for step in [
        "1. Start backend: mvn spring-boot:run -pl modules/api",
        "2. Start infrastructure: docker-compose up -d",
        "3. Start tunnel: cloudflared tunnel --url http://localhost:8080",
        "4. Load synthetic data: POST /api/v1/knowledge/load-synthetic",
        "5. Index to Pinecone: POST /api/v1/knowledge/records/{sysId}/reindex for each record",
        "6. Open ServiceNow Incident form and type in Short description",
        "7. AI resolution modal appears with suggested steps and code snippet",
        "8. (Optional) Show frontend dashboard at http://localhost:3000",
        "9. (Optional) Show Grafana dashboards at http://localhost:3001",
    ]:
        doc.add_paragraph(step, style="List Bullet")

    # ===== 16. ROADMAP =====
    doc.add_heading("16. Roadmap", level=1)
    for phase, desc in [
        ("Phase 1 - Current (POC)", "Core suggestion engine with Pinecone + Gemini, ServiceNow plugin, Docker Compose infrastructure"),
        ("Phase 2 - Production Hardening", "Live Gemini LLM integration, OAuth2 for ServiceNow, production tunnel (not LocalTunnel), real virus scanning"),
        ("Phase 3 - Enhanced UI", "Service Portal widget with rich card layout, suggestion acceptance/rejection tracking"),
        ("Phase 4 - Analytics & Feedback", "ROI analytics dashboard, agent feedback loop, deflection rate optimization"),
        ("Phase 5 - Multi-Tenant", "Support for multiple ServiceNow instances, workspace isolation, connector extensibility"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{phase}: ").bold = True
        p.add_run(desc)

    # -- Footer --
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("AI Service Desk Knowledge Intelligence Platform - Confidential  |  ").font.size = Pt(8)
    add_page_number(footer)

    return doc


if __name__ == "__main__":
    output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "output")
    os.makedirs(output_dir, exist_ok=True)
    doc = build_document()
    output_path = os.path.join(output_dir, "AI_Service_Desk_Architecture.docx")
    doc.save(output_path)
    print(f"Document saved: {output_path}")
